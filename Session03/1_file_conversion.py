"""
Slide 06 style — convert PDF / DOCX / TXT from data/raw into clean docs/*.txt
Run: python 1_file_conversion.py
"""
import os
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

RAW_DIR = Path("data/raw")
DOCS_DIR = Path("docs")


def clean_text(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def ensure_sample_docx():
    path = RAW_DIR / "sample_company_policy.docx"
    if path.exists():
        return path

    RAW_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Northstar Tech — Remote Work Policy", level=1)
    doc.add_paragraph(
        "Effective date: January 2025. All full-time employees may work remotely "
        "up to 3 days per week with manager approval."
    )
    doc.add_heading("Equipment", level=2)
    doc.add_paragraph(
        "The company provides a laptop and monitor for home office setup. "
        "Employees must use company-approved VPN when accessing internal systems."
    )
    doc.add_heading("Core Hours", level=2)
    doc.add_paragraph(
        "Employees must be available between 10:00 AM and 3:00 PM in their local timezone."
    )
    doc.add_heading("Expense Reimbursement", level=2)
    doc.add_paragraph(
        "Internet stipend of $50 per month is available for eligible remote employees."
    )
    doc.save(path)
    print(f"Created sample DOCX: {path}")
    return path


def read_txt(path):
    return path.read_text(encoding="utf-8")


def read_docx(path):
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(parts)


def read_pdf(path):
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text.strip())
    return "\n\n".join(pages)


def convert_one(source_path, output_dir):
    suffix = source_path.suffix.lower()
    output_path = output_dir / f"{source_path.stem}.txt"

    if suffix == ".txt":
        text = read_txt(source_path)
    elif suffix == ".docx":
        text = read_docx(source_path)
    elif suffix == ".pdf":
        text = read_pdf(source_path)
    else:
        print(f"Skip unsupported file: {source_path.name}")
        return None

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path.write_text(clean_text(text), encoding="utf-8")
    return output_path


def main():
    print("=== File Conversion: PDF / DOCX / TXT -> docs/*.txt ===\n")

    RAW_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    ensure_sample_docx()

    pdf_files = list(RAW_DIR.glob("*.pdf"))
    if not pdf_files:
        print("No PDF in data/raw/. Add data/raw/sample_document.pdf to demo PDF extraction.")
        print("Continuing with TXT and DOCX only.\n")

    sources = sorted(
        list(RAW_DIR.glob("*.txt"))
        + list(RAW_DIR.glob("*.docx"))
        + list(RAW_DIR.glob("*.pdf"))
    )

    if not sources:
        print("No source files found in data/raw/")
        return

    print(f"Found {len(sources)} source file(s):")
    for s in sources:
        print(f"  - {s.name}")

    print("\nConverting...")
    for source in sources:
        output = convert_one(source, DOCS_DIR)
        if output:
            chars = len(output.read_text(encoding="utf-8"))
            print(f"  {source.name} -> {output} ({chars} chars)")

    print("\nDone. Clean text is ready in docs/ for ingestion.")


if __name__ == "__main__":
    main()
